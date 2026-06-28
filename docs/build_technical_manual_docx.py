from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"E:\study\sensor-platform\docs\作物生长环境与表型关联分析系统技术手册.docx")

FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(35, 39, 47)
MUTED = RGBColor(92, 99, 112)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "CBD5E1"


def east_asia(run, font_name: str = FONT):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def para_format(paragraph, before=0, after=3, line=1.14):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, bold=False, size=10.5, color=TEXT):
    run = paragraph.add_run(text)
    east_asia(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    if level == 1:
        add_run(paragraph, text, bold=True, size=16, color=BLUE)
        para_format(paragraph, before=11, after=5, line=1.08)
    elif level == 2:
        add_run(paragraph, text, bold=True, size=13, color=BLUE)
        para_format(paragraph, before=8, after=4, line=1.08)
    else:
        add_run(paragraph, text, bold=True, size=12, color=DARK_BLUE)
        para_format(paragraph, before=6, after=3, line=1.08)


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    add_run(paragraph, text)
    para_format(paragraph)


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    add_run(paragraph, text)
    para_format(paragraph, after=1, line=1.12)


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    add_run(paragraph, text)
    para_format(paragraph, after=1, line=1.12)


def add_code(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    para_format(paragraph, before=1, after=3, line=1.0)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F8FAFC")
    p_pr.append(shd)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.7)
        run.font.color.rgb = RGBColor(30, 41, 59)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        rfonts.set(qn("w:eastAsia"), FONT)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color: str = BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].width = Cm(widths_cm[index])
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(header_cells[index], HEADER_FILL)
        borders(header_cells[index])
        p = header_cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, bold=True, size=9.2, color=DARK_BLUE)
        para_format(p, after=0, line=1.1)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Cm(widths_cm[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            borders(cells[index])
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, value, size=9.1)
            para_format(p, after=0, line=1.12)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, LIGHT_FILL)
    borders(cell, "D7DEE8")
    p = cell.paragraphs[0]
    add_run(p, f"{title}：", bold=True, color=DARK_BLUE)
    add_run(p, text)
    para_format(p, before=2, after=2, line=1.2)
    doc.add_paragraph()


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style.font.color.rgb = TEXT
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), FONT)
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)


def title_page(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "作物生长环境与表型关联分析系统", bold=True, size=24, color=RGBColor(11, 37, 69))
    para_format(title, after=8, line=1.1)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "技术手册", bold=True, size=18, color=BLUE)
    para_format(subtitle, after=18, line=1.1)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "当前系统版本 | 本地运行、状态检查、点云展示与迁移部署", size=11, color=MUTED)
    para_format(meta, after=24)
    add_callout(doc, "适用范围", "本文档以当前 sensor-platform 系统为准，覆盖采集模块、后端、前端、系统状态页、一键脚本和新电脑迁移部署。")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    configure(doc)
    title_page(doc)

    add_heading(doc, "1. 文档说明", 1)
    add_body(doc, "本文档以当前版本 sensor-platform 为准，覆盖系统组成、运行流程、页面功能、接口、数据库、迁移部署和一键启动/关闭脚本。")
    for item in [
        "本系统不负责传感器硬件数据采集算法本身。",
        "Java 采集模块负责接收传感器上报数据并写入 MySQL。",
        "Spring Boot 后端负责读取数据库、提供点云文件和系统状态接口。",
        "Vue 前端负责展示传感器趋势、点云、运行状态和交互控制。",
        "用户权限暂未实现，当前按本地电脑单机运行设计。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 系统组成", 1)
    add_table(doc, ["模块", "路径", "作用"], [
        ["Java 采集模块", "JavaSDKV2.2.2\\Demo", "监听 2404，接收传感器数据并写入 MySQL"],
        ["Spring Boot 后端", "backend", "提供 /api 接口，读取传感器数据、点云记录和点云文件"],
        ["Vue 前端", "frontend", "展示仪表盘、趋势图、点云和系统状态"],
        ["一键脚本", "start-system.bat / stop-system.bat / scripts", "启动和关闭本地运行环境"],
        ["文档", "docs", "技术手册、API 说明和 Word 手册"],
    ], [3.0, 4.9, 8.1])
    add_code(doc, """sensor-platform
├─ JavaSDKV2.2.2\\Demo
├─ backend
├─ frontend
├─ scripts
├─ docs
├─ start-system.bat
└─ stop-system.bat""")

    add_heading(doc, "3. 技术栈", 1)
    add_table(doc, ["层级", "技术", "当前用途"], [
        ["采集模块", "Java、RSNetDevice SDK、MySQL JDBC", "接收设备数据并写表"],
        ["后端服务", "Spring Boot、Maven、JDBC", "API 服务、数据库读取、本地文件读取、状态检查"],
        ["前端页面", "Vue 3、TypeScript、Vite", "仪表盘和交互界面"],
        ["图表", "ECharts", "温度、湿度、光照趋势图"],
        ["点云", "Three.js、PLYLoader、OrbitControls", ".ply 点云渲染和视角控制"],
        ["数据库", "MySQL", "保存传感器数据和点云索引"],
        ["启动脚本", "PowerShell、BAT", "本地一键启动和关闭"],
    ], [2.7, 5.4, 7.9])

    add_heading(doc, "4. 数据流", 1)
    add_body(doc, "传感器设备数据由 Java 采集模块写入 MySQL，Spring Boot 后端读取传感器表和点云索引表，并根据点云文件名从本地目录读取 .ply 文件，最终由 Vue 前端展示。")
    for item in [
        "传感器实时/存储数据写入 sensor_data_two。",
        "点云 .ply 文件由外部提供，文件本体放在本地目录。",
        "point_cloud.file_name 只保存文件名，后端根据配置目录查找文件。",
        "前端只访问 Spring Boot 后端接口，不直接连接数据库或读取本地文件。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 数据库设计", 1)
    add_heading(doc, "5.1 传感器表", 2)
    add_table(doc, ["字段", "类型建议", "说明"], [
        ["id", "BIGINT AUTO_INCREMENT", "可选主键，后端兼容没有该字段的情况"],
        ["device_id", "VARCHAR(64)", "设备编号"],
        ["node_id", "INT", "节点编号"],
        ["temperature", "DOUBLE", "温度"],
        ["humidity", "DOUBLE", "湿度"],
        ["light_intensity", "DOUBLE", "光照强度"],
        ["record_time", "DATETIME", "记录时间"],
    ], [3.0, 4.3, 8.7])
    add_code(doc, """CREATE TABLE IF NOT EXISTS sensor_data_two (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  device_id VARCHAR(64) NOT NULL,
  node_id INT NOT NULL,
  temperature DOUBLE,
  humidity DOUBLE,
  light_intensity DOUBLE,
  record_time DATETIME NOT NULL,
  INDEX idx_sensor_record_time (record_time)
);""")
    add_heading(doc, "5.2 点云表", 2)
    add_table(doc, ["字段", "类型建议", "说明"], [
        ["id", "BIGINT AUTO_INCREMENT", "主键"],
        ["record_time", "DATETIME", "点云记录时间"],
        ["file_name", "VARCHAR(255)", "点云文件名，例如 Stage_Flowering.ply"],
    ], [3.0, 4.3, 8.7])
    add_code(doc, """CREATE TABLE IF NOT EXISTS point_cloud (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  record_time DATETIME NOT NULL,
  file_name VARCHAR(255) NOT NULL,
  INDEX idx_point_cloud_record_time (record_time)
);""")

    add_heading(doc, "6. 后端配置", 1)
    add_body(doc, "后端配置文件位于 backend\\src\\main\\resources\\application.yml。迁移到新电脑时必须检查数据库地址、账号密码、点云目录和 CORS 前端地址。")
    add_code(doc, """server:
  port: 8080

app:
  datasource:
    url: jdbc:mysql://127.0.0.1:3306/sensor?useUnicode=true&characterEncoding=utf8&serverTimezone=Asia/Shanghai
    username: root
    password: "269756"
  point-cloud:
    directory: E:/study/new_sensor/pointCloud
  cors:
    allowed-origins: http://localhost:5173,http://127.0.0.1:5173""")

    add_heading(doc, "7. 采集模块说明", 1)
    for item in [
        "采集入口：JavaSDKV2.2.2\\Demo\\src\\main\\java\\demo\\Application.java。",
        "采集逻辑：JavaSDKV2.2.2\\Demo\\src\\main\\java\\demo\\DataCollector.java。",
        "监听端口：2404。",
        "依赖文件：param.dat 和 lib\\RSNetDevice-2.2.2.jar。",
        "写入字段：sensor_data_two(device_id, node_id, temperature, humidity, light_intensity, record_time)。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. 前端功能", 1)
    add_heading(doc, "8.1 仪表盘", 2)
    add_body(doc, "仪表盘包含顶部标题区、系统状态按钮、刷新按钮、概览指标、时间筛选、点云图、趋势图和右侧指标切换器。")
    add_heading(doc, "8.2 系统状态页", 2)
    add_table(doc, ["检查项", "检查方式", "说明"], [
        ["Spring Boot 后端", "当前接口响应", "后端能返回状态接口即为正常"],
        ["MySQL 数据库", "后端执行 SELECT 1", "检查数据库连接是否可用"],
        ["Java 采集模块", "后端检查 127.0.0.1:2404", "判断采集模块是否监听"],
        ["点云文件目录", "后端检查本地目录", "判断配置目录是否存在"],
    ], [3.5, 5.2, 7.3])
    add_code(doc, "GET /api/system/status")
    add_heading(doc, "8.3 趋势图和点云", 2)
    for item in [
        "趋势图使用 ECharts，支持综合视图、温度、湿度、光照、目标温度线和点云时间标记。",
        "点云使用 Three.js，支持 .ply 文件加载、轮播、视角控制、点大小、透明度、背景色和颜色模式。",
        "当前版本加入 ResizeObserver，图表和点云画布会随面板尺寸变化自动调整。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.4 当前 UI 调整", 2)
    for item in [
        "主仪表盘采用外层 flex + 内容区 CSS Grid，减少下方留白。",
        "右侧指标切换器已重做为侧边色条、圆形图标和右对齐数值的指标导航卡。",
        "点云控制面板支持收起和展开。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. API 接口", 1)
    add_table(doc, ["接口", "说明"], [
        ["GET /api/sensor/latest", "获取最新传感器数据"],
        ["GET /api/sensor/summary", "获取最新数据和总记录数"],
        ["GET /api/sensor/trend", "按 realtime/hour/day/week 查询趋势"],
        ["GET /api/point-cloud/list", "获取点云列表"],
        ["GET /api/point-cloud/latest", "获取最新点云"],
        ["GET /api/point-cloud/file/{fileName}", "读取点云文件"],
        ["GET /api/system/status", "检查后端、数据库、采集模块和点云目录状态"],
    ], [7.0, 9.0])

    add_heading(doc, "10. 启动和关闭", 1)
    add_heading(doc, "10.1 手动启动", 2)
    add_code(doc, """# 后端
Set-Location E:\\study\\sensor-platform\\backend
mvn spring-boot:run

# 前端
Set-Location E:\\study\\sensor-platform\\frontend
npm run dev

# 采集模块
Set-Location E:\\study\\sensor-platform\\JavaSDKV2.2.2\\Demo
mvn -q compile
mvn -q dependency:build-classpath "-Dmdep.outputFile=target/classpath.txt"
$cp = "target/classes;" + (Get-Content target/classpath.txt -Raw).Trim()
java -cp $cp demo.Application""")
    add_heading(doc, "10.2 一键启动和关闭", 2)
    add_table(doc, ["文件", "用途"], [
        ["start-system.bat", "双击启动本地系统"],
        ["stop-system.bat", "双击关闭本地系统"],
        ["scripts\\start-system.ps1", "实际启动逻辑，支持 -InstallDeps、-SkipCollector、-SkipMysqlStart"],
        ["scripts\\stop-system.ps1", "实际关闭逻辑，按 PID 文件和端口清理进程"],
    ], [5.2, 10.8])
    add_code(doc, """Set-Location E:\\study\\sensor-platform
.\\scripts\\start-system.ps1 -InstallDeps
.\\start-system.bat
.\\stop-system.bat
.\\scripts\\start-system.ps1 -SkipCollector""")

    add_heading(doc, "11. 功能验证", 1)
    add_code(doc, """# 构建
Set-Location E:\\study\\sensor-platform\\backend
mvn -q -DskipTests package
Set-Location E:\\study\\sensor-platform\\frontend
npm run build
Set-Location E:\\study\\sensor-platform\\JavaSDKV2.2.2\\Demo
mvn -q compile

# 端口
netstat -ano | Select-String ":8080|:5173|:2404"

# 接口
Invoke-RestMethod -Uri "http://localhost:8080/api/sensor/summary"
Invoke-RestMethod -Uri "http://localhost:8080/api/point-cloud/list"
Invoke-RestMethod -Uri "http://localhost:8080/api/system/status" """)
    for item in [
        "首页可进入系统。",
        "指标概览有数据。",
        "趋势图和点云图可显示。",
        "点云控制面板可收起和展开。",
        "右侧指标切换器可切换综合视图、温度、湿度、光照。",
        "系统状态页能正确显示后端、数据库、采集模块和点云目录状态。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. 新电脑迁移部署", 1)
    add_heading(doc, "12.1 需要拷贝的内容", 2)
    for item in [
        "sensor-platform 项目目录。",
        "JavaSDKV2.2.2\\Demo\\param.dat。",
        "JavaSDKV2.2.2\\Demo\\lib\\RSNetDevice-2.2.2.jar。",
        "点云 .ply 文件目录。",
        "如需历史数据，导出并迁移 MySQL 数据。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.2 基础软件", 2)
    add_table(doc, ["软件", "用途"], [
        ["JDK", "运行后端和采集模块"],
        ["Maven", "构建和启动 Java 项目"],
        ["Node.js / npm", "安装前端依赖和启动前端"],
        ["MySQL", "保存数据"],
        ["浏览器", "访问前端页面"],
    ], [4.0, 12.0])
    add_code(doc, """java -version
mvn -version
node -v
npm -v
mysql --version""")
    add_heading(doc, "12.3 数据库和配置", 2)
    add_code(doc, """CREATE DATABASE IF NOT EXISTS sensor DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;
mysqldump -uroot -p sensor > sensor_backup.sql
mysql -uroot -p sensor < sensor_backup.sql""")
    add_body(doc, "迁移后检查 application.yml 和 DataCollector.java 中的数据库地址、账号密码、点云目录、param.dat、RSNetDevice-2.2.2.jar 和 2404 端口。")
    add_heading(doc, "12.4 首次启动", 2)
    add_code(doc, """Set-Location E:\\study\\sensor-platform
.\\scripts\\start-system.ps1 -InstallDeps""")
    add_body(doc, "启动后打开 http://localhost:5173，再进入“系统状态”检查服务是否正常。")

    add_heading(doc, "13. 常见问题", 1)
    add_table(doc, ["问题", "排查方向"], [
        ["命令一直不结束", "后端、前端和采集模块都是常驻进程，端口监听后即为正常"],
        ["页面打不开", "检查 5173 是否监听，必要时重新启动前端"],
        ["数据加载失败", "检查 8080、MySQL、数据库账号密码、表结构和后端日志"],
        ["采集模块异常", "系统状态显示 2404 未监听时，需要启动 Java 采集模块"],
        ["点云不显示", "检查 file_name、点云目录、文件大小写和点云文件接口"],
        ["chunk 过大警告", "ECharts 和 Three.js 体积导致，不影响当前功能运行"],
    ], [4.2, 11.8])

    add_heading(doc, "14. 维护建议", 1)
    for item in [
        "数据库账号密码、点云目录建议后续迁移到环境变量。",
        "采集模块长期运行时建议改为 Windows 服务或计划任务。",
        "前端正式部署时可使用 npm run build 生成静态文件。",
        "点云文件较大时，可考虑分页、懒加载和文件大小提示。",
        "如果后续增加用户权限，建议在 Spring Boot 后端统一做认证和接口权限控制。",
    ]:
        add_bullet(doc, item)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(footer, "作物生长环境与表型关联分析系统技术手册", size=8, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
